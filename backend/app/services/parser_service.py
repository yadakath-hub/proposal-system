"""
Document parser — extract text from PDF, DOCX, XLSX files.
"""

import io
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


def detect_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".xlsx": "xlsx",
        ".xls": "xls",
    }
    return mapping.get(ext, "unknown")


def parse_pdf(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def parse_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def parse_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[openpyxl not installed — cannot parse xlsx]"

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sections = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"[{sheet}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sections)


def parse_file(data: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return parse_pdf(data)
    elif file_type in ("docx", "doc"):
        return parse_docx(data)
    elif file_type in ("xlsx", "xls"):
        return parse_xlsx(data)
    return ""
