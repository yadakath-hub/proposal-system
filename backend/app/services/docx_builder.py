"""
DOCX builder — assemble proposal sections into a formatted Word document.

Uses python-docx. Chinese font fallback: if 微軟正黑體 / 標楷體 are not
available on the host (common on Linux), the document will still render
correctly on Windows/macOS where the fonts exist.
"""

import io
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


DEFAULT_STYLES = {
    "title": {
        "font_name": "微軟正黑體",
        "font_size": 22,
        "bold": True,
        "alignment": "center",
    },
    "heading1": {
        "font_name": "微軟正黑體",
        "font_size": 16,
        "bold": True,
        "space_before": 12,
        "space_after": 6,
    },
    "heading2": {
        "font_name": "微軟正黑體",
        "font_size": 14,
        "bold": True,
        "space_before": 6,
        "space_after": 3,
    },
    "heading3": {
        "font_name": "微軟正黑體",
        "font_size": 12,
        "bold": True,
    },
    "body": {
        "font_name": "標楷體",
        "font_size": 12,
        "line_spacing": 1.5,
        "first_line_indent": 24,
    },
    "table": {
        "font_name": "標楷體",
        "font_size": 10,
    },
}


class DocxBuilder:
    def __init__(self, template_path: str | None = None, style_config: dict | None = None):
        if template_path:
            self.doc = Document(template_path)
        else:
            self.doc = Document()
        self.styles = {**DEFAULT_STYLES, **(style_config or {})}
        self._apply_default_styles()
        self._page_count_estimate = 0

    def _apply_default_styles(self) -> None:
        """Set sensible defaults for A4, margins, etc."""
        for section in self.doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    def _set_run_font(self, run, cfg: dict) -> None:
        font = run.font
        font_name = cfg.get("font_name", "標楷體")
        font.name = font_name
        # Set East Asian font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        font.size = Pt(cfg.get("font_size", 12))
        if cfg.get("bold"):
            font.bold = True
        if cfg.get("color"):
            font.color.rgb = RGBColor.from_string(cfg["color"])

    # -----------------------------------------------------------------
    # Cover page
    # -----------------------------------------------------------------

    def add_cover_page(
        self,
        project_name: str,
        company_name: str = "",
        tender_number: str = "",
        cover_date: str | None = None,
    ) -> None:
        # Add vertical space
        for _ in range(6):
            self.doc.add_paragraph("")

        # Project name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        self._set_run_font(run, {**self.styles.get("title", {}), "font_size": 26})

        self.doc.add_paragraph("")

        # Subtitle: 建議書
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("服務建議書")
        self._set_run_font(run, {**self.styles.get("title", {}), "font_size": 20})

        if tender_number:
            self.doc.add_paragraph("")
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"標案編號：{tender_number}")
            self._set_run_font(run, {**self.styles.get("body", {}), "font_size": 14})

        # Spacer
        for _ in range(4):
            self.doc.add_paragraph("")

        if company_name:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(company_name)
            self._set_run_font(run, {**self.styles.get("heading1", {}), "font_size": 18})

        # Date
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        display_date = cover_date or date.today().strftime("中華民國 %Y 年 %m 月")
        run = p.add_run(display_date)
        self._set_run_font(run, self.styles.get("body", {}))

        self.doc.add_page_break()
        self._page_count_estimate += 1

    # -----------------------------------------------------------------
    # Table of contents (TOC field — Word will update on open)
    # -----------------------------------------------------------------

    def add_table_of_contents(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("目  錄")
        self._set_run_font(run, self.styles.get("heading1", {}))

        # Insert a TOC field code
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = run2._element.makeelement(qn("w:instrText"), {})
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._element.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
        run3._element.append(fldChar2)

        run4 = paragraph.add_run("（請在 Word 中按 F9 更新目錄）")
        self._set_run_font(run4, self.styles.get("body", {}))

        run5 = paragraph.add_run()
        fldChar3 = run5._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run5._element.append(fldChar3)

        self.doc.add_page_break()
        self._page_count_estimate += 1

    # -----------------------------------------------------------------
    # Section content
    # -----------------------------------------------------------------

    def add_section(self, title: str, content: str, level: int = 1) -> None:
        heading_key = f"heading{min(level, 3)}"
        style_name = f"Heading {min(level, 3)}"

        heading = self.doc.add_heading(title, level=min(level, 3))
        for run in heading.runs:
            cfg = self.styles.get(heading_key, {})
            self._set_run_font(run, cfg)

        if content:
            for paragraph_text in content.split("\n"):
                text = paragraph_text.strip()
                if not text:
                    continue
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                body_cfg = self.styles.get("body", {})
                self._set_run_font(run, body_cfg)
                pf = p.paragraph_format
                if body_cfg.get("line_spacing"):
                    pf.line_spacing = body_cfg["line_spacing"]
                if body_cfg.get("first_line_indent"):
                    pf.first_line_indent = Pt(body_cfg["first_line_indent"])

        # Rough estimate: 1 page per ~2000 chars
        self._page_count_estimate += max(1, len(content) // 2000) if content else 0

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    # -----------------------------------------------------------------
    # Header / Footer
    # -----------------------------------------------------------------

    def set_header(self, text: str) -> None:
        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            if header.paragraphs:
                p = header.paragraphs[0]
            else:
                p = header.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def set_footer(self, text: str = "", include_page_number: bool = True) -> None:
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if text:
                run = p.add_run(text + "  ")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            if include_page_number:
                run = p.add_run()
                fldChar1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
                run._element.append(fldChar1)
                run2 = p.add_run()
                instrText = run2._element.makeelement(qn("w:instrText"), {})
                instrText.text = " PAGE "
                run2._element.append(instrText)
                run3 = p.add_run()
                fldChar2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
                run3._element.append(fldChar2)

    # -----------------------------------------------------------------
    # Table
    # -----------------------------------------------------------------

    def add_table(self, data: list[list[str]], headers: list[str] | None = None) -> None:
        cols = len(headers) if headers else (len(data[0]) if data else 0)
        if cols == 0:
            return

        table = self.doc.add_table(rows=0, cols=cols)
        table.style = "Table Grid"

        if headers:
            row = table.add_row()
            for i, h in enumerate(headers):
                cell = row.cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(self.styles.get("table", {}).get("font_size", 10))

        for row_data in data:
            row = table.add_row()
            for i, val in enumerate(row_data):
                if i < cols:
                    row.cells[i].text = str(val)

    # -----------------------------------------------------------------
    # Output
    # -----------------------------------------------------------------

    def save_to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.read()

    def save(self, output_path: str) -> None:
        self.doc.save(output_path)

    def get_page_count(self) -> int:
        return max(1, self._page_count_estimate)
